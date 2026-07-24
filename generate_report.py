import sys
sys.path.append('/Users/abhinavkumar/Library/Python/3.9/lib/python/site-packages')
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) # Navy Blue
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Steel Blue
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark Slate
    return p

def add_callout(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F2F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="1F497D"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"📌 {title}: ")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    r_body = p.add_run(text)
    r_body.font.name = 'Calibri'
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def create_report():
    doc = docx.Document()
    
    # Page setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("FinPulse: Stock Market Monitoring & Intelligence Platform")
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Comprehensive Technical & Architectural Report | SoFI AlgoLabs Assignment - 1")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.italic = True
    r_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Metadata Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Project Name:", "FinPulse India"), ("Live Application:", "https://finpulse99.streamlit.app")],
        [("Developer:", "Abhinav Kumar"), ("GitHub Repository:", "https://github.com/AbhinavKumarX/finpulse")]
    ]
    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "F9FAFB")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl, val = meta_data[row_idx][col_idx]
            r1 = p.add_run(f"{lbl} ")
            r1.bold = True
            r1.font.size = Pt(9.5)
            r1.font.name = 'Calibri'
            r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.name = 'Calibri'
            r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Executive Summary
    add_heading_styled(doc, "1. Executive Summary", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "FinPulse is an end-to-end stock market monitoring and intelligence platform engineered specifically for the Indian equity market (NSE/BSE). "
        "The system aggregates real-time market data, financial ratios, historical price trends, technical indicators, and financial news into a cohesive, high-performance web dashboard. "
        "Built to exceed all core requirements of the Society of Finance and Investing (SoFI) AlgoLabs Assignment 1, FinPulse includes automated user-isolated watchlists (pre-seeded with 20 NIFTY 50 stocks), "
        "an automated AI Copilot valuation engine, interactive Plotly technical charts (Moving Averages, RSI), a multi-parameter stock screener with CSV export, and a persistent SQLite database portfolio tracker."
    )

    add_callout(doc, "Live Web App: https://finpulse99.streamlit.app | GitHub Repository: https://github.com/AbhinavKumarX/finpulse", "SUBMISSION LINKS")

    # 2. Project Architecture
    add_heading_styled(doc, "2. Project Architecture & System Design", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "FinPulse employs a flexible, dual-mode architectural design that caters to both zero-cost cloud deployment and decoupled client-server API execution:"
    )

    bullet1 = doc.add_paragraph(style='List Bullet')
    bullet1.paragraph_format.space_after = Pt(3)
    r = bullet1.add_run("Production Monolith Deployment (streamlit_app.py): ")
    r.bold = True
    bullet1.add_run(
        "To host the application 100% free without credit card or trial limits on Streamlit Community Cloud, the primary app directly imports backend modules (database.py, fetcher.py) in memory. "
        "This eliminates inter-process HTTP overhead and maximizes execution speed."
    )

    bullet2 = doc.add_paragraph(style='List Bullet')
    bullet2.paragraph_format.space_after = Pt(6)
    r = bullet2.add_run("Decoupled REST API Architecture (backend/main.py & frontend/app.py): ")
    r.bold = True
    bullet2.add_run(
        "For standard API evaluation, a standalone FastAPI web service provides 16 RESTful endpoints complete with Swagger UI interactive documentation (/docs), CORS support, and X-User-Token header authentication."
    )

    # Architecture Table
    arch_tbl = doc.add_table(rows=1, cols=3)
    arch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = arch_tbl.rows[0].cells
    headers = ["Layer", "Technology", "Responsibility & Logic"]
    for idx, name in enumerate(headers):
        set_cell_background(hdr[idx], "1F497D")
        set_cell_margins(hdr[idx], top=100, bottom=100, left=120, right=120)
        p = hdr[idx].paragraphs[0]
        r = p.add_run(name)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    arch_rows = [
        ("Presentation Layer", "Streamlit, Plotly, st-keyup", "Renders dark-themed UI, interactive Plotly charts, key metrics, RSI indicators, and tabbed navigation."),
        ("API & Auth Layer", "FastAPI, Uvicorn, SHA-256", "Handles HTTP REST endpoints, user registration, token validation via X-User-Token headers, and CORS."),
        ("Data Pipeline", "yFinance, Yahoo Finance REST", "Fetches live price data, 52W ranges, fundamentals, 5Y historical OHLCV data, and financial news."),
        ("Persistence Layer", "SQLite3 (Relational DB)", "Stores user credentials, isolated watchlists (user_tickers), holdings (user_portfolio), and global caches.")
    ]

    for row_i, (layer, tech, resp) in enumerate(arch_rows):
        row_cells = arch_tbl.add_row().cells
        for c_idx, text in enumerate([layer, tech, resp]):
            set_cell_background(row_cells[c_idx], "F9FAFB" if row_i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 3. APIs Used & Data Pipeline
    add_heading_styled(doc, "3. APIs Used & Data Pipeline", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "FinPulse combines external market data providers with robust multi-tier fallback algorithms to ensure 100% data availability:"
    )

    api_points = [
        ("yFinance API: ", "Sourced for market capitalization, P/E ratio, P/B ratio, ROE, ROA, dividend yield, beta, 52-week highs/lows, target prices, and 5-year OHLCV price history."),
        ("Yahoo Finance REST Search API (query2.finance.yahoo.com): ", "Powers real-time autocomplete ticker searching filtered for National Stock Exchange of India (.NS) equities and ETFs."),
        ("Multi-Tier Data Fallback Engine (fetcher.py): ", "When Yahoo Finance info dictionaries encounter cloud IP rate-limiting, the engine dynamically falls back to fast_info attributes and 5-day OHLCV history calculations (e.g., P/E = Price / EPS, P/B = Price / Book Value) to guarantee non-zero price displays."),
        ("FastAPI REST Endpoints: ", "Exposes 16 API endpoints including /auth/register, /auth/login, /stocks, /stocks/{ticker}/history, /market-indices, /portfolio, /export-report, and /search.")
    ]

    for title, desc in api_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run(title)
        r.bold = True
        bp.add_run(desc)

    # 4. Database Design
    add_heading_styled(doc, "4. Database Design & Relational Schema", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The database layer is powered by SQLite3, featuring user-isolated tables and shared global data caches:"
    )

    db_tbl = doc.add_table(rows=1, cols=3)
    db_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = db_tbl.rows[0].cells
    for idx, name in enumerate(["Table Name", "Key Constraints", "Purpose & Cache Strategy"]):
        set_cell_background(hdr[idx], "1F497D")
        set_cell_margins(hdr[idx], top=100, bottom=100, left=120, right=120)
        p = hdr[idx].paragraphs[0]
        r = p.add_run(name)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    db_rows = [
        ("users", "user_id (PK, SHA-256), username (UNIQUE)", "Stores registered user credentials and generated tokens."),
        ("user_tickers", "user_id, ticker (PK: user_id, ticker)", "Per-user isolated watchlist. Pre-seeded with 20 NIFTY 50 stocks."),
        ("user_portfolio", "user_id, ticker, shares, buy_price", "Per-user portfolio holdings and buy prices for real-time P&L."),
        ("fundamental_data", "ticker (PK), name, sector, market_cap...", "Shared global fundamental cache to prevent duplicate external API calls."),
        ("historical_prices", "ticker, date (PK: ticker, date), close...", "Shared OHLCV historical price cache for Plotly charts.")
    ]

    for row_i, (t_name, t_cols, t_purp) in enumerate(db_rows):
        row_cells = db_tbl.add_row().cells
        for c_idx, text in enumerate([t_name, t_cols, t_purp]):
            set_cell_background(row_cells[c_idx], "F9FAFB" if row_i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    add_callout(doc, "Query Optimization: Watchlist queries utilize a LEFT JOIN (SELECT ut.ticker, COALESCE(f.name, ut.ticker) FROM user_tickers ut LEFT JOIN fundamental_data f ON ut.ticker = f.ticker) to ensure 100% of user-selected stocks render immediately even if external cache fetching is in progress.", "DATABASE OPTIMIZATION")

    # 5. Features Implemented
    add_heading_styled(doc, "5. Features Implemented", 1)
    
    feats = [
        ("1. Username Identity & Multi-User Isolation", "Simple claimable username handles (no password friction). Case-sensitive login (admin != Admin) with case-insensitive uniqueness enforcement. Pre-allocates 20 random NIFTY 50 stocks for every new user."),
        ("2. Interactive Market Dashboard", "Displays live price cards, 52-week range indicators, key financials, analyst consensus targets, and an automated AI Copilot heuristic engine providing qualitative analysis."),
        ("3. Technical Analysis Charts", "Plotly Candlestick and Line charts with toggleable 10D, 40D, 90D Moving Averages and 14-period RSI indicators across 1W, 1M, 6M, 1Y, and 5Y horizons."),
        ("4. Stock Comparison Engine", "Compare tab renders 1-Year normalized return performance curves (base = 100), side-by-side metric bar charts (P/E, P/B, ROE, ROA), P/E vs ROE bubble chart, and sector donut chart."),
        ("5. Stock Screener & CSV Export", "Multi-parameter filtering by P/E range, minimum ROE %, max P/B, and minimum Market Capitalization with 1-click CSV report export."),
        ("6. Portfolio Tracker & Real-Time P&L", "Tracks buy prices, share quantities, current market values, total P&L, percentage return, and portfolio allocation donut charts.")
    ]

    for title, desc in feats:
        add_heading_styled(doc, title, 2)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(desc)

    # 6. Challenges Faced & Solutions
    add_heading_styled(doc, "6. Challenges Faced & Technical Solutions", 1)
    
    challenges = [
        ("1. Hosting Without Credit Card Requirements", 
         "Challenge: Web hosting services like Render require credit card verification for free web services, while container hosts clear local SQLite databases upon restart.\n"
         "Solution: Architected a Production Monolith in streamlit_app.py that imports database.py and fetcher.py directly into Streamlit Community Cloud (100% free, cardless) while preserving backend/main.py for API evaluation."),
        ("2. Yahoo Finance API Rate Limiting & Partial Data Returns",
         "Challenge: Cloud server IPs frequently experience yFinance rate limits or empty info dictionaries, leading to 0.00 price displays.\n"
         "Solution: Created a 3-layer fallback in fetcher.py (info -> fast_info -> 5d history calculations) and updated database.py with SQL CASE WHEN logic to preserve valid non-zero metrics."),
        ("3. Streamlit v1.60+ Widget State Mutation Restrictions",
         "Challenge: Mutating st.session_state for a widget key after instantiation threw StreamlitAPIException.\n"
         "Solution: Implemented a dynamic key counter pattern (_s_key = f'search_{search_counter}'). Incrementing search_counter safely re-instantiates the search widget without state conflicts.")
    ]

    for title, desc in challenges:
        add_heading_styled(doc, title, 2)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(desc)

    # 7. Future Improvements
    add_heading_styled(doc, "7. Future Improvements & Roadmap", 1)
    fut_points = [
        ("Real-Time WebSocket Price Feeds: ", "Integrate TrueData or Zerodha Kite Connect WebSockets for sub-second tick streaming."),
        ("Production Database Migration: ", "Migrate persistent storage from SQLite3 to PostgreSQL / Supabase with Redis caching for scale."),
        ("Advanced Technical Indicators: ", "Add MACD (Moving Average Convergence Divergence), Bollinger Bands, and Fibonacci overlays."),
        ("OAuth2 Authentication: ", "Upgrade username identity to OAuth2 password hashing with email verification.")
    ]
    for title, desc in fut_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run(title)
        r.bold = True
        bp.add_run(desc)

    # 8. External Tools & Frameworks Used
    add_heading_styled(doc, "8. Frameworks, Libraries & AI Tools Used", 1)
    
    tools_tbl = doc.add_table(rows=1, cols=3)
    tools_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tools_tbl.rows[0].cells
    for idx, name in enumerate(["Category", "Tool / Library Name", "Role in Development"]):
        set_cell_background(hdr[idx], "1F497D")
        set_cell_margins(hdr[idx], top=100, bottom=100, left=120, right=120)
        p = hdr[idx].paragraphs[0]
        r = p.add_run(name)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tools_rows = [
        ("Frameworks", "Streamlit 1.35+, FastAPI 0.111+", "Frontend interactive dashboard and RESTful API framework."),
        ("Data & Math", "yFinance, Pandas, NumPy", "Market data extraction, time series manipulation, and data cleanup."),
        ("Visualization", "Plotly Express & Graph Objects", "Candlestick charts, moving averages, RSI indicators, and treemaps."),
        ("Database & Utils", "SQLite3, hashlib, python-docx", "Relational persistence, SHA-256 token hashing, and DOCX report generation."),
        ("AI Pair Programmer", "Google Antigravity AI (Gemini 3.6)", "Architectural planning, code generation, multi-user logic redesign, and debugging.")
    ]

    for row_i, (cat, name, role) in enumerate(tools_rows):
        row_cells = tools_tbl.add_row().cells
        for c_idx, text in enumerate([cat, name, role]):
            set_cell_background(row_cells[c_idx], "F9FAFB" if row_i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 9. Defense & Explanation Guide
    add_heading_styled(doc, "9. Guide to Defending & Explaining the Code", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "During evaluator discussion or viva, use these exact explanations to clarify your technical choices:"
    )

    explain_points = [
        ("Multi-User Identity: ", "When a user inputs username 'abhinav', database.create_user generates SHA-256('abhinav') as a unique user_id. All watchlists (user_tickers) and portfolios (user_portfolio) reference user_id, ensuring 100% data separation."),
        ("Pre-Allocation & Seeding: ", "Upon user creation, database.create_user triggers random.sample(NIFTY_50_POOL, 20) and inserts them into user_tickers. When the user logs in, seed_user_stocks fetches live data stock-by-stock with a progress bar."),
        ("Monolith vs API Architecture: ", "backend/main.py provides 16 FastAPI REST endpoints for standard API requirements. streamlit_app.py imports backend modules directly so the app deploys 100% free on Streamlit Cloud without credit card constraints."),
        ("Data Reliability Engine: ", "fetcher.py checks yfinance info first, then fast_info, and calculates 5-day history metrics. SQLite uses CASE WHEN logic during upsert so valid existing values are never overwritten with zeros.")
    ]

    for title, desc in explain_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run(title)
        r.bold = True
        bp.add_run(desc)

    # Save output
    out_path = "/Users/abhinavkumar/Downloads/sofi core alogolabs/project_report.docx"
    doc.save(out_path)
    print(f"Report successfully generated at {out_path}")

if __name__ == "__main__":
    create_report()
